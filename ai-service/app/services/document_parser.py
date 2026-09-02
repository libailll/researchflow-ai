from dataclasses import dataclass
from pathlib import Path
import re

import pymupdf
from docx import Document as WordDocument

from app.models.document import ParsedChunk


@dataclass(frozen=True)
class ParsedPage:
    page_number: int | None
    text: str


class DocumentParser:
    def __init__(self, chunk_size: int, chunk_overlap: int) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse_and_chunk(self, file_path: Path) -> list[ParsedChunk]:
        path = file_path.expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"文档文件不存在：{path}")
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            pages = self._parse_pdf(path)
        elif suffix == ".docx":
            pages = self._parse_docx(path)
        elif suffix in {".txt", ".md", ".markdown"}:
            pages = [ParsedPage(None, self._read_text(path))]
        else:
            raise ValueError(f"不支持的文件类型：{suffix}")

        chunks: list[ParsedChunk] = []
        for page in pages:
            for content in self._split_text(page.text):
                chunks.append(ParsedChunk(
                    pageNumber=page.page_number,
                    chunkIndex=len(chunks),
                    content=content,
                ))
        if not chunks:
            raise ValueError("文档中没有提取到可用文本，扫描版 PDF 需要后续接入 OCR")
        return chunks

    def _parse_pdf(self, path: Path) -> list[ParsedPage]:
        pages: list[ParsedPage] = []
        with pymupdf.open(path) as pdf:
            for index, page in enumerate(pdf):
                text = page.get_text("text", sort=True)
                if text and text.strip():
                    pages.append(ParsedPage(index + 1, text))
        return pages

    def _parse_docx(self, path: Path) -> list[ParsedPage]:
        document = WordDocument(path)
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return [ParsedPage(None, "\n\n".join(parts))]

    def _read_text(self, path: Path) -> str:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("文本文件编码无法识别，请转换为 UTF-8 后重试")

    def _split_text(self, text: str) -> list[str]:
        normalized = re.sub(r"[ \t]+", " ", text.replace("\r\n", "\n").replace("\r", "\n"))
        normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
        chunks: list[str] = []
        start = 0
        length = len(normalized)
        while start < length:
            end = min(start + self.chunk_size, length)
            if end < length:
                boundary_floor = start + int(self.chunk_size * 0.6)
                candidates = [normalized.rfind("\n", boundary_floor, end), normalized.rfind("。", boundary_floor, end), normalized.rfind(" ", boundary_floor, end)]
                boundary = max(candidates)
                if boundary > start:
                    end = boundary + 1
            content = normalized[start:end].strip()
            if content:
                chunks.append(content)
            if end >= length:
                break
            start = max(end - self.chunk_overlap, start + 1)
        return chunks
